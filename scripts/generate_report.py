#!/usr/bin/env python3
"""Generate the Decentralized Auction Platform Word report (.docx).

Run with the project venv:
    /Users/Cyberkid/decentralized-auction-platform/.venv/bin/python scripts/generate_report.py
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                   "Decentralized_Auction_Platform_Report.docx")

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)   # deep blue
ACCENT2 = RGBColor(0x2E, 0x86, 0xC1)  # lighter blue
GREEN = RGBColor(0x1E, 0x8A, 0x4C)


# ----------------------------------------------------------------------------- helpers

def add_field(paragraph, instr):
    """Insert a Word field (e.g. TOC, PAGE) into a paragraph."""
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_begin)

    instr_el = OxmlElement('w:instrText')
    instr_el.set(qn('xml:space'), 'preserve')
    instr_el.text = instr
    run._r.append(instr_el)

    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    run._r.append(fld_sep)

    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_end)


def set_cell_background(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def style_table(table, header=True):
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if header:
        for cell in table.rows[0].cells:
            set_cell_background(cell, '1F3A5F')
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def add_two_col_table(doc, header, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(header))
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].paragraphs[0].add_run(h)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            cells[i].paragraphs[0].add_run(str(val))
    style_table(table)
    if widths:
        for r in table.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.bold = bold
    r.font.italic = italic
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.add_run(text)
    return p


# ----------------------------------------------------------------------------- build

doc = Document()

# Base style tweaks
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

for lvl, color, size in [(1, ACCENT, 16), (2, ACCENT2, 13)]:
    st = doc.styles[f'Heading {lvl}']
    st.font.color.rgb = color
    st.font.size = Pt(size)
    st.font.bold = True

# ---- Footer with page number ----
section = doc.sections[0]
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.add_run("Page ")
add_field(fp, " PAGE ")
fp.add_run(" of ")
add_field(fp, " NUMPAGES ")

# ============================================================ TITLE PAGE
for _ in range(6):
    doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Decentralized Auction Platform")
r.font.size = Pt(30)
r.font.bold = True
r.font.color.rgb = ACCENT

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("A Smart-Contract Auction House on Ethereum")
r.font.size = Pt(16)
r.font.italic = True
r.font.color.rgb = ACCENT2

doc.add_paragraph()
tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tag.add_run("Where the contract is the auctioneer — no middleman, no trust required.")
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

for _ in range(8):
    doc.add_paragraph()

for label in ["Student / Team Name: ____________________________",
              "Course / Module: ____________________________",
              "Submission Date: 2026"]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = p.add_run(label)
    rr.font.size = Pt(12)

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ============================================================ TABLE OF CONTENTS
h1(doc, "Table of Contents")
para(doc, "This contents list is generated by Word. After opening the document, right-click "
          "anywhere on it and choose \"Update Field\" to fill in the page numbers.", italic=True)
toc_p = doc.add_paragraph()
add_field(toc_p, 'TOC \\o "1-2" \\h \\z \\u')
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ============================================================ 1. INTRODUCTION
h1(doc, "1. Introduction")
para(doc,
     "This project is a decentralized auction platform: an online auction house that runs "
     "entirely on the Ethereum blockchain. The central idea is simple but powerful — the smart "
     "contract IS the auctioneer. A smart contract is a small computer program that lives on the "
     "blockchain and runs exactly as written. Instead of trusting a person or a company to run "
     "the auction fairly, hold everyone's money, and pay out honestly, we trust a piece of public "
     "code that nobody can secretly change.")
para(doc,
     "In a traditional auction, an auctioneer takes the bids, decides who wins, collects the "
     "money, and hands it to the seller. You have to trust that person to be honest. In our system "
     "the contract does all of those jobs automatically: it lists items, records every bid, always "
     "knows the current highest bidder, closes the auction when time runs out, pays the seller, and "
     "returns money to everyone who lost. There is no middleman who can cheat, lose the money, or "
     "play favourites.")
para(doc, "Traditional Auction vs. This Decentralized Auction", bold=True)
add_two_col_table(doc,
    ["Question", "Traditional Auction", "This Decentralized Auction"],
    [
        ["Who runs it?", "A human auctioneer or company", "A smart contract (public code)"],
        ["Who holds the money?", "The auction house's bank account", "The contract, locked by code"],
        ["How much trust is needed?", "You must trust the auctioneer", "You trust open, unchangeable code"],
        ["Can bids be hidden or faked?", "Possibly — you can't see everything", "No — every bid is public and recorded"],
        ["How is the seller paid?", "Manually, later, by the company", "Automatically, the moment it closes"],
        ["Middleman / fees?", "Yes — commission to the house", "No middleman; only small network fees"],
        ["Transparency", "Limited", "Full — anyone can verify every action"],
        ["Hours of operation", "Business hours / set sale dates", "24 hours a day, 7 days a week"],
    ],
    widths=[1.6, 2.3, 2.6])

# ============================================================ 2. HOW IT WORKS
h1(doc, "2. How It Works — The Five Steps")
para(doc, "The whole platform can be understood as five plain-English steps. The contract performs "
          "every one of them automatically.")
h2(doc, "Step 1 — List")
para(doc, "A seller creates an auction (called a \"lot\") for an item. They give it a name, a "
          "description, a starting price, and a length of time it should stay open. Each new lot gets "
          "its own number: lot #1, lot #2, and so on, so many auctions can run at the same time.")
h2(doc, "Step 2 — Bid")
para(doc, "Buyers place bids by sending ETH (Ethereum's digital money) to the contract. The amount "
          "of ETH you send IS your bid. The very first bid must be at least the starting price, and "
          "every later bid must be strictly higher than the current highest bid. The contract safely "
          "holds the money the whole time.")
h2(doc, "Step 3 — Track")
para(doc, "At all times the contract knows the current highest bid and who made it. When someone is "
          "outbid, the contract makes a note that their money is waiting to be collected, so nobody's "
          "funds are ever lost.")
h2(doc, "Step 4 — End")
para(doc, "Once the timer runs out, the auction can be closed. Importantly, anyone can press the "
          "\"close\" button — it does not have to be the seller — so the auction can never get stuck "
          "if the seller disappears. The contract refuses to close early.")
h2(doc, "Step 5 — Pay")
para(doc, "When the auction closes, the winning bid is sent automatically to the seller. Everyone "
          "who was outbid can then withdraw their money back to their own wallet whenever they like. "
          "No one has to ask permission or wait for a human to process a refund.")

# ============================================================ 3. ARCHITECTURE
h1(doc, "3. System Architecture")
para(doc, "The system has three parts that work together. You can think of them as the shop window, "
          "the cashier's keycard, and the vault.")
h2(doc, "The Frontend (the website)")
para(doc, "A single web page (index.html) is the part people see and click. It shows the list of "
          "auctions, the live countdown timers, and the buttons for bidding, closing, and withdrawing. "
          "It is just an ordinary web page — it does not hold any money itself.")
h2(doc, "The Wallet (MetaMask)")
para(doc, "MetaMask is a free browser extension that acts as the user's digital wallet and identity. "
          "When the website needs to send a bid or any other action to the blockchain, MetaMask pops "
          "up and asks the user to approve and sign it. Nothing happens with the user's money without "
          "their explicit approval in MetaMask.")
h2(doc, "The Smart Contract (on Ethereum)")
para(doc, "The contract is the program living on the Ethereum blockchain. It enforces all the rules "
          "and is the only thing that actually holds the bids. It is the trustworthy vault and "
          "rule-keeper at the centre of everything.")
h2(doc, "How a bid travels through the system")
para(doc, "When a user places a bid, the information flows like this:")
para(doc, "User clicks \"Place Bid\" on the website  →  the website asks MetaMask to sign the "
          "transaction  →  the user approves in MetaMask  →  the transaction is sent to the smart "
          "contract on the Ethereum blockchain  →  the contract checks the rules and records the new "
          "highest bid  →  the website reads the updated information back from the blockchain and "
          "refreshes the auction card for everyone to see.", italic=True)

# ============================================================ 4. FUNCTIONS
h1(doc, "4. Function-by-Function Breakdown")
para(doc, "The contract is made up of named actions called functions. The table below explains, in "
          "plain English, what each one does.")
para(doc, "Action functions (these change something on the blockchain)", bold=True)
add_two_col_table(doc,
    ["Function", "What it does in plain English"],
    [
        ["createAuction", "Lists a new item for sale. The seller provides a name, description, "
                          "starting price and duration; the contract opens a new numbered lot."],
        ["placeBid", "Places a bid. The ETH attached to the call is the bid. It must beat the "
                     "starting price (first bid) or the current highest bid (later bids). The "
                     "previous leader is credited so they can get their money back. The seller is "
                     "not allowed to bid on their own item."],
        ["endAuction", "Closes the auction after its deadline (and only once). It sends the winning "
                       "bid to the seller. Anyone may call it, so the auction is never stuck."],
        ["withdrawRefund", "Lets a bidder who was outbid pull their money back out of the contract, "
                           "safely and at any time."],
    ],
    widths=[1.6, 4.9])

para(doc, "View functions (these only read information and are free to use)", bold=True)
add_two_col_table(doc,
    ["Function", "What it tells you"],
    [
        ["getAuction", "All the details of one lot: seller, item, prices, highest bidder, end time, "
                       "and whether it has closed."],
        ["getBidHistory", "The complete list of every bid placed on a lot, for full transparency."],
        ["bidCount", "How many bids a lot has received."],
        ["timeLeft", "How many seconds remain before a lot closes (used for the countdown)."],
        ["myRefund", "How much money the person asking is owed from a lot."],
    ],
    widths=[1.6, 4.9])

para(doc, "Events (announcements the blockchain broadcasts)", bold=True)
para(doc, "Whenever something important happens, the contract emits an \"event\" — a public "
          "announcement that the website can listen for. The four events are:")
bullet(doc, "AuctionCreated — a new lot has been listed.")
bullet(doc, "BidPlaced — someone placed a bid.")
bullet(doc, "AuctionEnded — a lot has closed and the seller has been paid.")
bullet(doc, "RefundWithdrawn — an outbid bidder has reclaimed their money.")

# ============================================================ 5. CONCEPTS
h1(doc, "5. Key Blockchain Concepts")
para(doc, "This section explains the handful of technical ideas behind the project in everyday terms.")
h2(doc, "Smart contracts")
para(doc, "A smart contract is a program stored on the blockchain. Once it is published, it runs "
          "exactly as written and cannot be quietly altered. It can hold money and move it according "
          "to its rules, which is what makes a trustless auction possible.")
h2(doc, "Wallets")
para(doc, "A wallet (such as MetaMask) is an app that stores your digital keys and lets you approve "
          "actions on the blockchain. It is both your identity and your way of signing off on "
          "transactions.")
h2(doc, "ETH vs. wei")
para(doc, "ETH is the currency of Ethereum. Wei is its smallest unit, a bit like pennies are to a "
          "pound, except there are a lot more of them: 1 ETH = 1,000,000,000,000,000,000 wei "
          "(that is 10 to the power of 18). The contract always works in wei to avoid rounding "
          "errors, and the website converts to and from friendly ETH amounts for the user.")
h2(doc, "Gas")
para(doc, "Every action that changes the blockchain costs a small fee called gas, which pays the "
          "network for the computing work. Reading information is free; sending a bid or closing an "
          "auction costs a little gas.")
h2(doc, "Testnets")
para(doc, "A testnet is a practice version of Ethereum that uses pretend money with no real value. "
          "It behaves just like the real network, so it is perfect for learning and for school "
          "projects — you can do everything for free. This project uses the Sepolia testnet.")

# ============================================================ 6. FRONTEND
h1(doc, "6. The Frontend (index.html)")
para(doc, "The user interface is a single, self-contained web page. It needs no installation and no "
          "build step — you simply open the file in a browser. Its features are:")
bullet(doc, "Connect MetaMask, then display the connected account and which network it is on.")
bullet(doc, "A field to paste the deployed contract's address. The address is remembered between "
            "visits using the browser's localStorage, so you do not have to paste it every time.")
bullet(doc, "A create-auction form with fields for the name, description, starting price in ETH, "
            "and a duration with a minutes / hours / days selector.")
bullet(doc, "Each auction is shown as a card with the item, description, seller, current highest "
            "bid and bidder, a live countdown that ticks every second, and a colour-coded status "
            "pill (live, ending soon, awaiting close, or closed).")
bullet(doc, "Buttons to place a bid, close the auction once its deadline passes, and withdraw a "
            "refund when the user is owed money. The user's own address is highlighted.")
bullet(doc, "An expandable bid history for each lot, showing every bid that was placed.")
bullet(doc, "Human-readable error messages: if the user rejects a transaction, has insufficient "
            "funds, or breaks a rule, the page shows a clear sentence rather than a confusing code.")

# ============================================================ 7. SECURITY
h1(doc, "7. Security Decisions (and Why)")
para(doc, "Because the contract handles real value, it is written defensively. Each decision below "
          "exists to stop a specific problem.")
h2(doc, "Pull-over-push refunds")
para(doc, "The contract never tries to automatically send money back to people who were outbid. "
          "Instead it records what each person is owed and lets them withdraw it themselves. Why? "
          "Imagine a malicious bidder whose wallet is programmed to reject incoming money. If the "
          "contract tried to force a refund to them, the payment would fail — and that failure could "
          "freeze the entire auction. By letting each person \"pull\" their own refund, one bad actor "
          "can only ever inconvenience themselves.")
h2(doc, "Checks-effects-interactions (blocking reentrancy)")
para(doc, "When paying out a refund, the contract first sets the person's recorded balance to zero, "
          "and only then sends the money. This order matters. A known attack called \"reentrancy\" "
          "tries to call the withdraw function again in the middle of a payment to drain extra funds. "
          "Because the balance is already zero by then, the attacker gets nothing extra.")
h2(doc, "The seller cannot bid")
para(doc, "The contract rejects any bid from the seller on their own item, preventing them from "
          "secretly inflating the price.")
h2(doc, "Deadline enforcement")
para(doc, "Bids are refused after the deadline, and the auction cannot be closed before it. The "
          "clock is enforced by the blockchain, not by a person.")
h2(doc, "Anyone can close the auction")
para(doc, "Closing is open to anyone after the deadline, so the auction always reaches its "
          "conclusion even if the seller loses interest or disappears.")

# ============================================================ 8. BUILD GUIDE
h1(doc, "8. Free Build-and-Test Guide")
para(doc, "Everything in this project can be built and tested for free. There are two paths, and we "
          "recommend doing them in order. Both the Remix VM and the Sepolia testnet are 100% FREE — "
          "they never use real money.", bold=True)
h2(doc, "Path A — Remix VM (start here, completely free)")
para(doc, "The Remix VM is a pretend blockchain that runs inside your web browser. It is the fastest "
          "way to see the contract work, and it costs nothing.")
numbered(doc, "Open a web browser and go to remix.ethereum.org.")
numbered(doc, "In the file panel, create a new file named AuctionPlatform.sol.")
numbered(doc, "Paste in the contract code from contracts/AuctionPlatform.sol.")
numbered(doc, "Open the \"Solidity Compiler\" tab and compile using version 0.8.20.")
numbered(doc, "Open the \"Deploy & Run Transactions\" tab and set the Environment to \"Remix VM\".")
numbered(doc, "Click \"Deploy\". The contract appears at the bottom with buttons for every function.")
numbered(doc, "Use the buttons to create an auction, place bids (type a value in the VALUE field), "
              "fast-forward is not needed because you control the test accounts, then end the "
              "auction and withdraw. All for free.")
h2(doc, "Path B — Sepolia testnet with MetaMask (still completely free)")
para(doc, "Once it works in the Remix VM, you can put it on a real public test network. Sepolia uses "
          "free test ETH, so this also costs nothing.")
numbered(doc, "Install the MetaMask browser extension and create a wallet.")
numbered(doc, "In MetaMask, switch the network to \"Sepolia\".")
numbered(doc, "Get free test ETH from a Sepolia faucet (search \"Sepolia faucet\" and paste in your "
              "wallet address). This fake ETH has no real-world value.")
numbered(doc, "Back in Remix, set the Environment to \"Injected Provider - MetaMask\" and deploy. "
              "Approve the transaction in MetaMask.")
numbered(doc, "Copy the deployed contract's address from Remix.")
numbered(doc, "Open frontend/index.html in your browser, paste the address into the address field, "
              "and click Connect.")
numbered(doc, "Now create auctions, place bids, close auctions and withdraw refunds through the "
              "website — every action is confirmed in MetaMask and recorded on Sepolia, all for free.")

# ============================================================ 9. VERIFICATION
h1(doc, "9. Verification — Automated Tests")
para(doc, "An automated test script runs a complete auction from start to finish on a local test "
          "blockchain and checks that the contract behaves correctly at every step. Running "
          "\"npx hardhat test\" reports 2 passing tests, covering all of the required checks below.")
add_two_col_table(doc,
    ["#", "What is tested", "Result"],
    [
        ["1", "Create an auction", "PASS"],
        ["2", "Reject a bid below the starting price", "PASS"],
        ["3", "Reject the seller bidding on their own item", "PASS"],
        ["4", "Bidder A bids, then Bidder B outbids", "PASS"],
        ["5", "Reject a bid that is not higher than the current highest", "PASS"],
        ["6", "A's pending refund equals their original bid", "PASS"],
        ["7", "Reject ending the auction before the deadline", "PASS"],
        ["8", "The bid history contains both bids", "PASS"],
        ["9", "After the deadline, end the auction; seller receives exactly the winning bid", "PASS"],
        ["10", "A withdraws their refund successfully", "PASS"],
        ["11", "A second withdrawal is rejected", "PASS"],
    ],
    widths=[0.5, 5.0, 1.0])

# ============================================================ 10. EXTENSIONS
h1(doc, "10. Possible Extensions")
para(doc, "The platform is deliberately focused, but it could be extended in several directions:")
bullet(doc, "Minimum bid increments — require each bid to beat the last by a set amount.")
bullet(doc, "Anti-sniping — automatically extend the timer when a bid arrives in the final seconds.")
bullet(doc, "NFT items (ERC-721) — auction unique digital collectibles and transfer them to the winner.")
bullet(doc, "Reserve prices — let sellers set a hidden minimum below which the item will not sell.")
bullet(doc, "Token bidding (ERC-20) — allow bids in a stablecoin or other token instead of ETH.")
bullet(doc, "Item images via IPFS — attach pictures stored on decentralized file storage.")

# ============================================================ 11. TEAM
h1(doc, "11. Team Division of Work")
para(doc, "The work was divided across the following roles (member names are placeholders to fill in).")
add_two_col_table(doc,
    ["Role", "Member", "Main Responsibilities"],
    [
        ["Smart Contract Developer", "________", "Wrote and commented AuctionPlatform.sol; "
                                                 "designed the auction logic and security."],
        ["Frontend Developer", "________", "Built index.html; MetaMask connection, auction cards, "
                                           "countdowns and error handling."],
        ["Tester / QA", "________", "Wrote the Hardhat test scenario; verified all 11 checks pass."],
        ["Report & Documentation", "________", "Produced this report and the README quickstart guide."],
    ],
    widths=[2.1, 1.2, 3.2])

# ============================================================ 12. CONCLUSION
h1(doc, "12. Conclusion")
para(doc, "This project delivers a working decentralized auction platform in which the smart contract "
          "acts as the auctioneer. It lists items, accepts and tracks bids, closes auctions on time, "
          "pays sellers automatically, and refunds losing bidders — all without a trusted middleman. "
          "The contract is written defensively, with pull-over-push refunds and the "
          "checks-effects-interactions pattern guarding against well-known attacks. A self-contained "
          "web page provides a clean interface through MetaMask, and an automated test suite "
          "confirms the rules are enforced correctly. Best of all, the whole system can be built and "
          "tested entirely for free using the Remix VM and the Sepolia testnet, making it ideal as a "
          "demonstration of how trust can be replaced by transparent, verifiable code.")

# ============================================================ 13. GLOSSARY
h1(doc, "13. Glossary")
add_two_col_table(doc,
    ["Term", "Meaning"],
    [
        ["Blockchain", "A shared, tamper-resistant public ledger of transactions maintained by many computers."],
        ["Ethereum", "A blockchain that can run programs (smart contracts), not just record payments."],
        ["Smart Contract", "A program stored on the blockchain that runs exactly as written and can hold and move money."],
        ["Solidity", "The programming language used to write Ethereum smart contracts."],
        ["Wallet", "An app that stores your keys and lets you approve blockchain actions."],
        ["MetaMask", "A popular browser-extension wallet used to connect to Ethereum websites."],
        ["ETH", "Ether — the digital currency of Ethereum."],
        ["Wei", "The smallest unit of ETH. 1 ETH = 10^18 wei."],
        ["Gas", "A measure of the computing work a transaction needs."],
        ["Gas Fee", "The small payment to the network for performing a transaction."],
        ["Testnet", "A free practice blockchain that uses valueless pretend money."],
        ["Sepolia", "The Ethereum testnet used by this project."],
        ["Mainnet", "The real Ethereum network, where ETH has real value."],
        ["Transaction", "A signed instruction sent to the blockchain (e.g. placing a bid)."],
        ["Address", "A unique identifier for a wallet or contract, like an account number."],
        ["Bid", "An offer of ETH for an item in an auction."],
        ["Auctioneer", "The party that runs an auction. Here, that party is the smart contract."],
        ["Pull-over-push", "A safety pattern where users withdraw money owed to them rather than being sent it automatically."],
        ["Reentrancy", "An attack that re-enters a function mid-payment to steal funds; blocked here by careful ordering."],
        ["Remix", "A free in-browser tool for writing, compiling and deploying smart contracts."],
        ["ethers.js", "A JavaScript library the website uses to talk to the blockchain."],
        ["localStorage", "A small browser storage area used here to remember the contract address."],
    ],
    widths=[1.6, 4.9])

doc.save(OUT)
print("Saved:", OUT)
